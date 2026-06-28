"""
ImpactFrame — Streamlit Web App
No CSS. Pure native Streamlit + .streamlit/config.toml for branding.
"""

import streamlit as st
import json
import sys
import os
import time
from pathlib import Path
from dotenv import load_dotenv
st.markdown("""
<link href="https://fonts.googleapis.com/css2?family=Libre+Baskerville:wght@700&display=swap" rel="stylesheet">
<style>
    details > summary { font-size: 16px !important; font-weight: 600 !important; color: #1E2D4E !important; }
    details > summary p { font-size: 16px !important; font-weight: 600 !important; }
</style>
""", unsafe_allow_html=True)

load_dotenv()

sys.path.insert(0, str(Path(__file__).parent))
from adk_pipeline import run_adk_pipeline as run_pipeline

# ── Asset helper ─────────────────────────────────────────────────
ASSETS = Path(__file__).parent / "assets"

def asset(name):
    p = ASSETS / name
    return str(p) if p.exists() else None

# ── Page config ──────────────────────────────────────────────────
st.set_page_config(
    page_title="ImpactFrame",
    page_icon=asset("icon.png") or "📊",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Helpers ──────────────────────────────────────────────────────
def fmt_money(n):
    try: return f"${int(n):,}"
    except: return str(n)

def fmt_pct(n):
    try: return f"{float(n):.1f}%"
    except: return str(n)

def dir_icon(d):
    d = (d or "MAINTAIN").upper()
    if d == "INCREASE": return "↑"
    if d == "DECREASE": return "↓"
    return "→"

def sev_icon(s):
    return {"HIGH": "🔴", "MEDIUM": "🟡", "LOW": "🔵"}.get(s, "🔵")

def grade_icon(g):
    return {"A": "🟢", "B": "🔵", "C": "🟡", "D": "🔴"}.get(g, "⚪")

def section_header(icon_file, title):
    """Render a section header with icon beside title."""
    ico = asset(icon_file)
    if ico:
        c1, c2 = st.columns([0.05, 0.95])
        with c1:
            st.image(ico, width=54)
        with c2:
            st.markdown(
                f"<div style='font-size:18px;font-weight:700;color:#1E2D4E;padding-top:10px;'>{title}</div>",
                unsafe_allow_html=True
            )
    else:
        st.markdown(
            f"<div style='font-size:18px;font-weight:700;color:#1E2D4E;'>{title}</div>",
            unsafe_allow_html=True
        )

# ── PDF + DOCX export helpers ────────────────────────────────────
def generate_pdf(data: dict) -> bytes:
    """Generate a simple PDF report from results."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors
        import io

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4,
                                rightMargin=50, leftMargin=50,
                                topMargin=60, bottomMargin=60)

        NAVY   = HexColor("#1E2D4E")
        ORANGE = HexColor("#E87722")
        LIGHT  = HexColor("#F5F5F5")

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("title", fontSize=20, textColor=NAVY,
                                     spaceAfter=6, fontName="Helvetica-Bold")
        h2_style = ParagraphStyle("h2", fontSize=14, textColor=NAVY,
                                  spaceBefore=12, spaceAfter=6, fontName="Helvetica-Bold")
        h3_style = ParagraphStyle("h3", fontSize=12, textColor=ORANGE,
                                  spaceBefore=8, spaceAfter=4, fontName="Helvetica-Bold")
        body_style = ParagraphStyle("body", fontSize=10, textColor=HexColor("#222222"),
                                    spaceAfter=4, fontName="Helvetica")
        caption_style = ParagraphStyle("caption", fontSize=9, textColor=HexColor("#666666"),
                                       spaceAfter=4, fontName="Helvetica-Oblique")

        alloc     = data.get("allocation", {})
        conflicts = data.get("conflicts", {})
        scores    = data.get("scores", {})
        recs      = alloc.get("allocation_recommendations", {})
        clist     = conflicts.get("conflicts", [])

        story = []

        # Title
        story.append(Paragraph("ImpactFrame", title_style))
        story.append(Paragraph("Resource Allocation Intelligence Report", h2_style))
        story.append(Spacer(1, 12))

        # Executive Summary
        story.append(Paragraph("Executive Summary", h2_style))
        story.append(Paragraph(alloc.get("allocation_summary", ""), body_style))
        story.append(Spacer(1, 10))

        # Stats
        total    = fmt_money(alloc.get("total_budget", 0))
        n_conf   = conflicts.get("conflict_count", len(clist))
        n_review = sum(1 for v in recs.values() if v.get("human_review_required"))

        stats_data = [
            ["Total Budget", "Programs", "Conflicts", "Need Review"],
            [total, str(len(recs)), str(n_conf), str(n_review)],
        ]
        t = Table(stats_data, colWidths=[120, 100, 100, 100])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), NAVY),
            ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
            ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",   (0, 0), (-1, -1), 10),
            ("ALIGN",      (0, 0), (-1, -1), "CENTER"),
            ("BACKGROUND", (0, 1), (-1, 1), LIGHT),
            ("GRID",       (0, 0), (-1, -1), 0.5, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [LIGHT, colors.white]),
        ]))
        story.append(t)
        story.append(Spacer(1, 14))

        # Allocations
        story.append(Paragraph("Budget Allocation Recommendations", h2_style))
        for prog, rec in recs.items():
            conf_pct = int(float(rec.get("confidence", 0)) * 100)
            story.append(Paragraph(
                f" {prog} — {fmt_money(rec.get('recommended_amount', 0))} "
                f"({fmt_pct(rec.get('recommended_percentage', 0))})", h3_style))
            story.append(Paragraph(rec.get("primary_rationale", ""), body_style))
            if rec.get("human_review_required"):
                story.append(Paragraph(
                    rec.get("human_review_reason", ""),
                    caption_style))
            story.append(Spacer(1, 6))

        # Conflicts
        story.append(Paragraph("Data Conflicts", h2_style))
        for c in clist:
            sev = c.get("severity", "LOW")
            story.append(Paragraph(
                f"{sev_icon(sev)} [{sev}] {c.get('program_area', '')} — {c.get('description', '')}",
                body_style))
            story.append(Paragraph(
                f"Action: {c.get('recommended_human_action', '')}", caption_style))
            story.append(Spacer(1, 4))

        # Priorities
        priorities = alloc.get("implementation_priorities", [])
        if priorities:
            story.append(Paragraph("Implementation Priorities", h2_style))
            for p in priorities:
                story.append(Paragraph(
                    f"{p.get('priority', '?')}. {p.get('action', '')} "
                    f"— {p.get('program', '')} · {p.get('timeline', '')}",
                    body_style))

        # Board flags
        flags = alloc.get("flags_for_board", [])
        if flags:
            story.append(Paragraph("Items for Board Decision", h2_style))
            for f in flags:
                story.append(Paragraph(f"⚠ {f}", body_style))

        doc.build(story)
        return buffer.getvalue()

    except ImportError:
        return None


def generate_docx(data: dict) -> bytes:
    """Generate a Word document report from results."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io

        alloc     = data.get("allocation", {})
        conflicts = data.get("conflicts", {})
        recs      = alloc.get("allocation_recommendations", {})
        clist     = conflicts.get("conflicts", [])

        doc = Document()

        # Title
        title = doc.add_heading("ImpactFrame", 0)
        title.runs[0].font.color.rgb = RGBColor(0x1E, 0x2D, 0x4E)

        doc.add_heading("Resource Allocation Intelligence Report", 1)

        # Summary
        doc.add_heading("Executive Summary", 2)
        doc.add_paragraph(alloc.get("allocation_summary", ""))

        # Allocations
        doc.add_heading("Budget Allocation Recommendations", 2)
        for prog, rec in recs.items():
            d = (rec.get("change_direction") or "MAINTAIN").upper()
            h = doc.add_heading(
                f"{dir_icon(d)} {prog} — {fmt_money(rec.get('recommended_amount', 0))} "
                f"({fmt_pct(rec.get('recommended_percentage', 0))})", 3)
            doc.add_paragraph(rec.get("primary_rationale", ""))
            if rec.get("human_review_required"):
                p = doc.add_paragraph(
                    f"Human Review Required: {rec.get('human_review_reason', '')}")
                p.runs[0].italic = True

        # Conflicts
        doc.add_heading("Data Conflicts", 2)
        for c in clist:
            sev = c.get("severity", "LOW")
            doc.add_heading(
                f"{sev_icon(sev)} [{sev}] {c.get('program_area', '')} "
                f"— {c.get('description', '')}", 3)
            doc.add_paragraph(f"Action: {c.get('recommended_human_action', '')}")

        # Priorities
        priorities = alloc.get("implementation_priorities", [])
        if priorities:
            doc.add_heading("Implementation Priorities", 2)
            for p in priorities:
                doc.add_paragraph(
                    f"{p.get('priority', '?')}. {p.get('action', '')} "
                    f"— {p.get('program', '')} · {p.get('timeline', '')}",
                    style="List Number")

        # Board flags
        flags = alloc.get("flags_for_board", [])
        if flags:
            doc.add_heading("Items for Board Decision", 2)
            for f in flags:
                doc.add_paragraph(f"⚠ {f}", style="List Bullet")

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    except ImportError:
        return None


# ════════════════════════════════════════════════════════════════
# SIDEBAR
# ════════════════════════════════════════════════════════════════
with st.sidebar:

    ico = asset("icon.png")
    if ico:
        st.image(ico, width=100)
    st.markdown("### ImpactFrame")
    st.caption("Where Policy, Data and Resources Make Impact")
    st.divider()

    # Data sources
    st.subheader("Data Sources")
    sources = [
        ("policy.png",    "Policy Reports"),
        ("data.png",      "Health Data"),
        ("budget.png",    "Budget"),
        ("surveys.png",   "Surveys"),
        ("resources.png", "Resources"),
    ]
    for icon_file, label in sources:
        ico = asset(icon_file)
        c1, c2 = st.columns([1, 4])
        with c1:
            if ico:
                st.image(ico, width=42)
        with c2:
            st.caption(label)

    st.divider()

    # Agent pipeline
    st.subheader("Agent Pipeline")
    agents = [
        ("1", "Source Agent",         "Collects all data via MCP"),
        ("2", "Evidence Agent",       "Extracts key facts"),
        ("3", "Conflict Agent",       "Flags disagreements"),
        ("4", "Confidence Agent",     "Scores reliability"),
        ("5", "Recommendation Agent", "Decision support"),
    ]
    for num, name, desc in agents:
        st.markdown(f"**{num}. {name}**")
        st.caption(desc)

    st.divider()

    # Pipeline run results
    if "pipeline_meta" in st.session_state and st.session_state.pipeline_meta:
        meta = st.session_state.pipeline_meta
        st.markdown("##### 🏥 Riverside Community Health Clinic")
        st.caption("Serving 12,000 patients annually")
        st.divider()
        st.markdown(f"**Total Budget Reviewed**")
        st.markdown(f"### {meta.get('total_budget', '—')}")
        st.caption(f"{meta.get('program_count', 0)} programs evaluated")
        st.divider()
        st.markdown(f"**Pipeline Run Time**")
        st.markdown(f"### :orange[~{meta.get('run_time', '?')} seconds]")
        st.caption(f"5 agents · {meta.get('api_calls', 5)} API calls")
        st.divider()
        step_times = meta.get("step_times", {})
        for i, name in enumerate(["Source Agent","Evidence Agent","Conflict Agent","Confidence Agent","Allocation Agent"], 1):
            t = f" · {step_times[name]}s" if name in step_times else ""
            st.markdown(f"{i}. **{name}**")
            st.caption(f"✓ complete{t}")
        st.divider()

    # API status
    api_key = os.getenv("GOOGLE_API_KEY", "")
    if api_key and api_key != "your_google_api_key_here":
        st.success("✅ Gemini Connected")
    else:
        st.error("❌ AI Engine Not Connected")
        st.caption("Add GOOGLE_API_KEY to your .env file")

    st.caption("Kaggle AI Agents Capstone · Agents for Good")


# ════════════════════════════════════════════════════════════════
# MAIN HEADER
# ════════════════════════════════════════════════════════════════
# ═══════════════════════════════════════════════════════
# HEADER
# ═══════════════════════════════════════════════════════

logo = asset("logo.png")

left, right = st.columns([1.3, 2])

with left:
    if logo:
        st.image(logo, width=550)

with right:

    st.markdown("""
        <div style="
            font-family:'Libre Baskerville', serif;
            font-size:48px;
            font-weight:700;
            line-height:1.1;
            margin-bottom:18px;
        ">
            <span style="color:#1E2D4E;">Impact</span><span style="color:#E87722;">Frame</span>
        </div>
    """, unsafe_allow_html=True)

    st.markdown(
        """
        <div style="
            font-size:25px;
            font-weight:600;
            color:#1E2D4E;
            margin-top:18px;
            line-height:1.2;
        ">
            From Evidence to Action
        </div>

        <div style="
            font-size:16px;
            color:#757575;
            margin-top:14px;
            line-height:1.5;
        ">
            Transparent recommendations for smarter resource allocation.
        </div>

        <div style="
            font-size:15px;
            font-weight:600;
            color:#E87722;
            margin-top:30px;
            margin-bottom:15px;
            letter-spacing:0.3px;
        ">
            Where Policy, Data and Resources Make Impact
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.write("")

    

# --- Run controls -----------------------------

title_col, btn_col1, btn_col2 = st.columns(
    [3, 1.2, 1.2],
    vertical_alignment="center"
)
with title_col:
    st.markdown("""
    <div style="
        font-size:22px;
        font-weight:500;
        color:#1E2D4E;
        margin-top:10px;
    ">
        Start Analysis
    </div>
    """, unsafe_allow_html=True)

with btn_col1:
    run_btn = st.button(
    
        "▶ Run Analysis ",
        type="primary",
        use_container_width=True
    )

with btn_col2:
    demo_btn = st.button(
        " 📂 Load Demo",
        use_container_width=True
    )
st.markdown("""
<div style="
font-size:17px;
color:#5F6368;
margin-top:10px;
margin-bottom:20px;
">
Upload your data or load the demo dataset to generate transparent,
evidence-based funding recommendations.
</div>
""", unsafe_allow_html=True)

st.divider()

# ── Session state ─────────────────────────────────────────────────
if "results" not in st.session_state:
    st.session_state.results = None

if "error" not in st.session_state:
    st.session_state.error = None



# ── Run pipeline ──────────────────────────────────────────────────
if run_btn:
    if not api_key or api_key == "your_google_api_key_here":
        st.error("❌ AI Engine not connected. Add GOOGLE_API_KEY to your .env file.")
    else:
        st.session_state.results = None
        st.session_state.error   = None

        progress_bar = st.progress(0, text="Starting pipeline...")
        step_cols    = st.columns(5)
        placeholders = []
        for i, col in enumerate(step_cols):
            with col:
                p = col.empty()
                p.markdown(f"⬜ Agent {i+1}")
                placeholders.append(p)

        agent_names = [
            "Source Agent",
            "Evidence Agent",
            "Conflict Agent",
            "Confidence Agent",
            "Recommendation Agent",
        ]

        def on_progress(step, msg):
            pct = int((step / 5) * 100)
            progress_bar.progress(pct, text=msg)
            for i, p in enumerate(placeholders):
                name = agent_names[i]
                if i + 1 < step:
                    p.markdown(f"✅ {name}")
                elif i + 1 == step:
                    p.markdown(f"🔄 **{name}**")
                else:
                    p.markdown(f"⬜ {name}")

        try:
            with st.spinner("Running pipeline..."):
                results = run_pipeline(progress_callback=on_progress)

            progress_bar.progress(100, text="Complete!")

            output_path = Path(__file__).parent / "output" / "latest_results.json"
            output_path.parent.mkdir(exist_ok=True)
            with open(output_path, "w") as f:
                json.dump(results, f, indent=2)

            alloc_tmp = results.get("allocation", {})
            recs_tmp  = alloc_tmp.get("allocation_recommendations", {})
            st.session_state.pipeline_meta = {
                "run_time":      results.get("pipeline_run_time_seconds", "—"),
                "api_calls":     results.get("api_calls", 5),
                "step_times":    results.get("step_times", {}),
                "total_budget":  fmt_money(alloc_tmp.get("total_budget", 0)),
                "program_count": len(recs_tmp),
            }
            st.session_state.results = results
            st.success("✅ Analysis complete!")
            time.sleep(1)
            st.rerun()

        except Exception as e:
            st.session_state.error = str(e)
            st.error(f"❌ {e}")
            import traceback
            with st.expander("Error details"):
                st.code(traceback.format_exc())

# ── Load demo ─────────────────────────────────────────────────────
if demo_btn:
    demo_path = Path(__file__).parent / "output" / "demo_results.json"
    if demo_path.exists():
        with open(demo_path) as f:
            st.session_state.results = json.load(f)
        st.session_state.pipeline_meta = None
        st.rerun()
    else:
        st.warning("No demo file found. Run the pipeline once first.")


# ════════════════════════════════════════════════════════════════
# RESULTS
# ════════════════════════════════════════════════════════════════
if st.session_state.results:
    r         = st.session_state.results

    _src_labels = {
        "get_health_data":           "Population Health Data",
        "get_budget_constraints":    "Budget & Finance Records",
        "get_survey_results":        "Community Survey Results",
        "get_program_effectiveness": "Program Effectiveness Report",
        "population_summary":        "Population Health Data",
        "budget_summary":            "Budget & Finance Records",
        "survey_summary":            "Community Survey Results",
        "effectiveness_summary":     "Program Effectiveness Report",
        "staff_summary":             "Staff Availability Data",
    }
    def friendly_src(s):
        return _src_labels.get(s, s.replace("_", " ").replace("get ", "").title())
    alloc     = r.get("allocation", {})
    conflicts = r.get("conflicts", {})
    scores    = r.get("scores", {})
    recs      = alloc.get("allocation_recommendations", {})
    clist     = conflicts.get("conflicts", [])
    n_conf    = conflicts.get("conflict_count", len(clist))
    n_review  = sum(1 for v in recs.values() if v.get("human_review_required"))

    # ── Always visible stats row ─────────────────────────────────
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.metric("Total Budget",    fmt_money(alloc.get("total_budget", 0)))
    with c2:
        st.metric("Programs",        len(recs))
    with c3:
        st.metric("Conflicts",       n_conf)
    with c4:
        st.metric("Need Review",     n_review)

    st.divider()

    # ── Executive Summary — always visible ───────────────────────
    section_header("impact.png", "Executive Summary")
    st.info(alloc.get("allocation_summary", "Analysis complete."))

    st.divider()

    with st.expander("💰 Budget Review", expanded=True):
        st.caption("Evidence-based funding recommendations across all priority programs.")
        st.write("")

        for prog, rec in recs.items():

            d = (rec.get("change_direction") or "MAINTAIN").upper()
            icon = dir_icon(d)
            conf_pct = int(float(rec.get("confidence", 0)) * 100)

            change = rec.get("recommended_amount", 0) - rec.get("current_amount", 0)
            change_abs = abs(int(change))

            sign = "+" if change >= 0 else "-"
            color = "green" if change >= 0 else "red"
            label = "Funding Adjustment"

            with st.expander(f"{prog} • {fmt_money(rec.get('recommended_amount',0))}"):

                # ---------------- Metrics ----------------
                ca, cb, cc = st.columns(3)

                with ca:
                    st.metric(
                        "Recommended Budget",
                        fmt_money(rec.get("recommended_amount", 0))
                    )

                    st.markdown(
                        """
                        <div style="
                            color:#5F6368;
                            font-size:14px;
                            font-weight:500;
                            margin-top:8px;
                            margin-bottom:4px;
                        ">
                            Funding Adjustment
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )   

                    st.markdown(
                        f"""
                        <div style="
                            color:#1E2D4E;
                            font-size:20px;
                            font-weight:500;
                        ">
                            {sign}{fmt_money(change_abs)}
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )
                    


                with cb:
                    st.metric(
                        "Existing Allocation",
                        fmt_money(rec.get("current_amount", 0))
                    )

                with cc:
                    if conf_pct >= 80:
                        level = "High"
                    elif conf_pct >= 60:
                        level = "Moderate"
                    else:
                        level = "Low"

                    st.markdown("**Data Confidence Score**")

                    st.markdown(
                        f"""
                        <div style="font-size:28px;
                                    font-weight:400;
                                    color:#1E2D4E;
                                    line-height:1.1;">
                            {level}
                        </div>

                        <div style="
                            font-size:18px;
                            color:#6B7280;
                            margin-top:6px;
                        ">
                            {conf_pct}%
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )                   




# ---------- Rationale ----------
                st.markdown(
                    f"**Rationale**\n\n{rec.get('primary_rationale', '')}"
                )
                evidence = rec.get("evidence_used", [])

                if evidence:
                    with st.expander("📊 Evidence Summary"):
                        for ev in evidence:
                            st.markdown(f"- {ev}")
                st.divider()

                if rec.get("human_review_required"):
                    st.info(
                        f"👤 **Human Review Recommendation**\n\n"
                        f"{rec.get('human_review_reason', '')}"
                    )
    st.divider()

    # ── Conflicts — expandable, HIGH auto-open ───────────────────
    with st.expander("Data Conflicts", expanded=False):
        section_header("conflicts.png", f"Data Conflicts ({n_conf})")

        if not clist:
            st.success("No conflicts detected between data sources.")
        else:
            for c in clist:
                sev      = c.get("severity", "LOW")
                auto_open = (sev == "HIGH")

                with st.expander(
                    f"{sev_icon(sev)} {c.get('conflict_id', '')} · "
                    f"{c.get('program_area', '')} · {sev}",
                    expanded=auto_open
                ):
                    st.markdown(f"**{c.get('description', '')}**")

                    ca, cb = st.columns(2)
                    with ca:
                        st.markdown(f"**Source:** {friendly_src(c.get('source_a', ''))}")
                        st.caption(c.get("source_a_claim", ""))
                    with cb:
                        st.markdown(f"**Source:** {friendly_src(c.get('source_b', ''))}")
                        st.caption(c.get("source_b_claim", ""))

                    st.warning(f"👤 **Action Required:** {c.get('recommended_human_action', '')}")

    st.divider()

    # ── Confidence Scores — expandable ───────────────────────────
    with st.expander("Confidence Scores", expanded=False):
        section_header("confidence.png", "Confidence Scores")
        st.caption("Grades reflect how reliable each data source is for making budget decisions. A = Very reliable → D = Low reliability")

        src_conf  = scores.get("source_confidence", {})
        prog_conf = scores.get("program_confidence", {})

        if src_conf:
            st.markdown("**Source Reliability**")

            src_labels = {
                "get_health_data": "Health Data",
                "get_budget_constraints": "Budget",
                "get_survey_results": "Surveys",
                "get_program_effectiveness": "Program Reports",
            }

            cols = st.columns(len(src_conf))
            for i, (src, sc) in enumerate(src_conf.items()):
                with cols[i]:
                    score_pct = int(float(sc.get("score", 0)) * 100)
                    grade     = sc.get("grade", "B")
                    st.markdown(f"### {grade} {grade_icon(grade)}")
                    st.markdown(
                        f"<span style='background:#f0f0f0;padding:2px 8px;"
                        f"border-radius:10px;font-size:13px;color:#444'>"
                        f"Score: {score_pct}%</span>",
                        unsafe_allow_html=True
                    )
                    st.caption(sc.get("rationale", ""))

        if prog_conf:
            st.markdown("**Program Evidence Strength**")
            for prog, pc in prog_conf.items():
                grade    = pc.get("grade", "B")
                strength = pc.get("evidence_strength", "MODERATE")
                score    = int(float(pc.get("overall_score", 0)) * 100)
                ca, cb, cc, cd = st.columns([3, 1, 1, 3])
                with ca: st.write(prog)
                with cb: st.write(f"**{grade}** {grade_icon(grade)}")
                with cc: st.caption(strength)
                with cd: st.progress(score / 100)

    st.divider()

    # ── Implementation Priorities — expandable ───────────────────
    priorities = alloc.get("implementation_priorities", [])
    if priorities:
        with st.expander("Implementation Priorities", expanded=False):
            section_header("priority.png", "Implementation Priorities")
            for p in priorities:
                ca, cb = st.columns([0.08, 0.92])
                with ca:
                    st.markdown(
                        f"<div style='background:#E87722;color:white;"
                        f"border-radius:50%;width:30px;height:30px;"
                        f"display:flex;align-items:center;justify-content:center;"
                        f"font-weight:bold;font-size:14px'>{p.get('priority','?')}</div>",
                        unsafe_allow_html=True
                    )
                with cb:
                    st.markdown(f"**{p.get('action', '')}**")
                    st.caption(f"{p.get('program', '')} · {p.get('timeline', '')}")

        st.divider()

    # ── Key Tradeoffs — expandable ───────────────────────────────
    tradeoffs = alloc.get("key_tradeoffs", [])
    if tradeoffs:
        with st.expander("Key Tradeoffs", expanded=False):
            section_header("tradeoffs.png", "Key Tradeoffs")
            for t in tradeoffs:
                st.markdown(f"- {t}")

        st.divider()

    # ── Board Decisions — expandable ─────────────────────────────
    flags = alloc.get("flags_for_board", [])
    if flags:
        with st.expander("Items for Board Decision", expanded=False):
            section_header("board.png", "Items for Board Decision")
            for f in flags:
                st.warning(f"⚠️ {f}")

        st.divider()

    # ── Download section ─────────────────────────────────────────
    st.markdown("#### Download Report")

    ca, cb, cc = st.columns([1, 1, 2])

    with ca:
        pdf_bytes = generate_pdf(r)
        if pdf_bytes:
            st.download_button(
                label="⬇ PDF Report",
                data=pdf_bytes,
                file_name="impactframe_report.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
        else:
            st.caption("Install reportlab for PDF export")

    with cb:
        docx_bytes = generate_docx(r)
        if docx_bytes:
            st.download_button(
                label="⬇ Word Report",
                data=docx_bytes,
                file_name="impactframe_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        else:
            st.caption("Install python-docx for Word export")

    with cc:
        st.download_button(
            label="⬇ JSON (Raw Data)",
            data=json.dumps(r, indent=2),
            file_name="impactframe_report.json",
            mime="application/json",
            use_container_width=True,
        )


elif st.session_state.error:
    st.error(f"❌ {st.session_state.error}")

else:

    # Data source icons
    items = [
        ("policy.png",    "Policy Reports"),
        ("data.png",      "Health Data"),
        ("budget.png",    "Budget"),
        ("surveys.png",   "Surveys"),
        ("resources.png", "Resources"),
    ]
    cols = st.columns(5)
    for col, (icon_f, label) in zip(cols, items):
        with col:
            ico = asset(icon_f)
            if ico:
                import base64
                img_b64 = base64.b64encode(open(ico, "rb").read()).decode()
                st.markdown(
                    f"""<div style="text-align:center;padding:4px 4px;">
                        <img src="data:image/png;base64,{img_b64}"
                             style="width:180px;height:180px;object-fit:contain;">
                    </div>""",
                    unsafe_allow_html=True
                )

    st.divider()
    st.markdown("#### What happens when you run:")

    agents_info = [
        ("1", "Source Agent",         "Collects all data from 4 MCP tools"),
        ("2", "Evidence Agent",       "Extracts key facts per program area"),
        ("3", "Conflict Agent",       "Detects where sources disagree"),
        ("4", "Confidence Agent",     "Scores reliability of each source"),
        ("5", "Recommendation Agent", "Generates evidence-based decision support"),
    ]
    for num, name, desc in agents_info:
        ca, cb = st.columns([0.08, 0.92])
        with ca:
            st.markdown(
                f"<div style='background:#E87722;color:white;"
                f"border-radius:50%;width:30px;height:30px;"
                f"display:flex;align-items:center;justify-content:center;"
                f"font-weight:bold;font-size:14px'>{num}</div>",
                unsafe_allow_html=True
            )
        with cb:
            st.markdown(f"**{name}** — {desc}")
